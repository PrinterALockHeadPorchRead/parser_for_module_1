from docx import Document
import xml.etree.ElementTree as ET
from docx.exceptions import InvalidFileFormatError
from typing import List, Dict, Optional, Union

class DOCXProcessor:
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
        self.doc = self._load_document()
        self.is_valid = self._validate_syntax()
        self.text_content = self._extract_text()
        self.tables = self._extract_tables()
        self.metadata = self._extract_metadata()

    def _load_document(self) -> Optional[Document]:
        """Загрузка документа"""
        try:
            return Document(self.file_path)
        except InvalidFileFormatError:
            print(f"Ошибка: файл {self.file_path} не является валидным DOCX")
            return None
            
        except FileNotFoundError:
            print(f"Ошибка: файл {self.file_path} не найден")
            return None
            
        except PermissionError:
            print(f"Ошибка доступа: недостаточно прав для чтения файла")
            return None
        
        except Exception as e:
            print(f"Непредвиденная ошибка загрузки: {str(e)}")
            return None

    def _validate_syntax(self) -> bool:
        """Проверка XML-структуры документа"""
        if not self.doc:
            return False
        try:
            xml_str = self.doc._element.xml
            ET.fromstring(xml_str)
            return True
        except ET.ParseError as e:
            print(f"XML ошибка: {e}")
            return False

    def _extract_text(self) -> str:
        """Извлечение текста"""
        if not self.is_valid:
            return ""
            
        try:
            return "\n".join(para.text for para in self.doc.paragraphs)
            
        except AttributeError:
            print("Ошибка: документ не содержит параграфов")
            return ""
        
        except Exception as e:
            print(f"Непредвиденная ошибка извлечения текста: {str(e)}")
            return ""

    def _extract_tables(self) -> List[List[List[str]]]:
        """Извлечение таблиц"""
        if not self.is_valid:
            return []
            
        tables = []
        try:
            for table in self.doc.tables:
                rows = []
                for row in table.rows:
                    try:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(cells)
                    except (AttributeError, IndexError):
                        print("Пропущена поврежденная строка таблицы")
                        continue
                tables.append(rows)
            return tables
            
        except AttributeError:
            print("Ошибка: неожиданная структура таблиц")
            return []
        except Exception as e:
            print(f"Непредвиденная ошибка таблиц: {str(e)}")
            return []

    def _extract_metadata(self) -> Dict[str, Union[str, int]]:
        """Получение метаданных"""
        if not self.is_valid:
            return {}
        try:
            core = self.doc.core_properties
            return {
                "author": core.author,
                "created": core.created,
                "modified": core.modified
            }
        except AttributeError:
            print("Предупреждение: метаданные отсутствуют")
            return {}
        except Exception as e:
            print(f"Ошибка доступа к метаданным: {str(e)}")
            return {}

    def print_results(self) -> None:
        print(f"XML-структура: {'Валидна' if self.is_valid else 'Ошибка'}")
        
        print("\nТекст документа:")
        print(self.text_content[:500] + "\n..." if len(self.text_content) > 500 else self.text_content)
        
        print("\nТаблицы из документа:")
        for i, table in enumerate(self.tables, 1):
            print(f"Таблица {i}:")
            for row in table:
                print(" | ".join(row))
                
        print("\nМетаданные:")
        for key, value in self.metadata.items():
            print(f"{key}: {value}")