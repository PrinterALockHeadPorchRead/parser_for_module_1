import pytest
from pathlib import Path
from datetime import datetime
from parsers.parser_docx import DOCXProcessor
from docx import Document

@pytest.fixture
def create_valid_docx(tmp_path):
    """Создает временный валидный DOCX-файл."""
    file = tmp_path / "valid.docx"
    doc = Document()
    doc.add_paragraph("Первый параграф.")
    doc.add_paragraph("Второй параграф.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Ячейка 1"
    table.cell(0, 1).text = "Ячейка 2"
    table.cell(1, 0).text = "Ячейка 3"
    table.cell(1, 1).text = "Ячейка 4"
    doc.core_properties.author = "Test Author"
    doc.core_properties.created = datetime(2023, 1, 1)
    doc.core_properties.modified = datetime(2023, 1, 2)
    doc.save(str(file))
    return str(file)

@pytest.fixture
def create_invalid_docx(tmp_path):
    """Создает временный невалидный DOCX-файл."""
    file = tmp_path / "invalid.docx"
    with open(file, "wb") as f:
        f.write(b"This is not a valid DOCX")
    return str(file)

# Тест на загрузку валидного документа
def test_load_valid_document(create_valid_docx):
    processor = DOCXProcessor(create_valid_docx)
    assert processor.doc is not None
    assert processor.is_valid is True

# Тест на загрузку невалидного документа
def test_load_invalid_document(create_invalid_docx):
    processor = DOCXProcessor(create_invalid_docx)
    assert processor.doc is None
    assert processor.is_valid is False

# Тест на проверку XML-структуры
def test_validate_syntax_valid(create_valid_docx):
    processor = DOCXProcessor(create_valid_docx)
    assert processor._validate_syntax() is True

def test_validate_syntax_invalid(create_invalid_docx):
    processor = DOCXProcessor(create_invalid_docx)
    assert processor._validate_syntax() is False

# Тест на извлечение текста
def test_extract_text_valid(create_valid_docx):
    processor = DOCXProcessor(create_valid_docx)
    expected_text = "Первый параграф.\nВторой параграф."
    assert processor.text_content == expected_text

def test_extract_text_invalid(create_invalid_docx):
    processor = DOCXProcessor(create_invalid_docx)
    assert processor.text_content == ""

# Тест на извлечение таблиц
def test_extract_tables_valid(create_valid_docx):
    processor = DOCXProcessor(create_valid_docx)
    expected_tables = [
        [["Ячейка 1", "Ячейка 2"], ["Ячейка 3", "Ячейка 4"]]
    ]
    assert processor.tables == expected_tables

def test_extract_tables_invalid(create_invalid_docx):
    processor = DOCXProcessor(create_invalid_docx)
    assert processor.tables == []

# Тест на извлечение метаданных
def test_extract_metadata_valid(create_valid_docx):
    """Проверка, что метаданные извлекаются корректно."""
    processor = DOCXProcessor(create_valid_docx)
    expected_metadata = {
        "author": "Test Author",
        "created": datetime(2023, 1, 1),  # Без временной зоны
        "modified": datetime(2023, 1, 2),  # Без временной зоны
    }
    # Удаляем временную зону из фактических данных
    actual_metadata = {
        "author": processor.metadata["author"],
        "created": processor.metadata["created"].replace(tzinfo=None),
        "modified": processor.metadata["modified"].replace(tzinfo=None),
    }
    assert actual_metadata == expected_metadata

def test_extract_metadata_invalid(create_invalid_docx):
    processor = DOCXProcessor(create_invalid_docx)
    assert processor.metadata == {}

# Тест на вывод результатов
def test_print_results_valid(create_valid_docx, capsys):
    processor = DOCXProcessor(create_valid_docx)
    processor.print_results()
    captured = capsys.readouterr()
    assert "XML-структура: Валидна" in captured.out
    assert "Текст документа:" in captured.out
    assert "Первый параграф." in captured.out
    assert "Таблица 1:" in captured.out
    assert "Ячейка 1 | Ячейка 2" in captured.out
    assert "Метаданные:" in captured.out
    assert "author: Test Author" in captured.out

def test_print_results_invalid(create_invalid_docx, capsys):
    processor = DOCXProcessor(create_invalid_docx)
    processor.print_results()
    captured = capsys.readouterr()
    assert "XML-структура: Ошибка" in captured.out
    assert "Текст документа:" in captured.out
    assert "Таблицы из документа:" in captured.out
    assert "Метаданные:" in captured.out
